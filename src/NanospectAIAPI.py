import logging
import os
import sys
import subprocess
from pathlib import Path

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from PIL import Image
try:
    from pypdf import PdfReader, PdfWriter
except Exception:
    try:
        from PyPDF2 import PdfReader, PdfWriter
    except Exception:
        raise ImportError("Neither 'pypdf' nor 'PyPDF2' is available; install one (e.g. pip install pypdf).")

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    docx2pdf_convert = None
    DOCX2PDF_AVAILABLE = False

import google.generativeai as genai


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

# Configurations
INPUT_DOCX = "testforJackson.docx"
OUTPUT_DOCX = "output_images_captions_summaries.docx"
OUTPUT_PDF = "inspection.pdf"
COVER_PDF = "cover.pdf"
FINAL_PDF = "final_report.pdf"
IMAGE_TEMP_DIR = "extracted_images"

MODEL_NAME = "gemini-2.5-pro"  

# Ensure Gemini API key present
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise RuntimeError("GEMINI_API_KEY environment variable is not set. Set it before running this script.")
genai.configure(api_key=api_key)
model = genai.GenerativeModel(MODEL_NAME)

# Namespaces used for parsing docx images
XML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}


# Cover page generation
def create_cover(output_path: str):
    """Generate the cover page PDF interactively (prompts user)."""
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    margin = 0.75 * inch
    logo_h = 1.1 * inch
    logo_w = 1.3 * inch
    #Draw SAU Computer Science Logo at top left
    c.drawImage("saulogo.jpg", margin, height - margin - logo_h, width=logo_w, height=logo_h, preserveAspectRatio=True, mask="auto")
    c.setStrokeColorRGB(0.2, 0.2, 0.2)

    x_text = width - margin - 0.25 * inch
    y = height - margin - 0.2 * inch

    # Project Title
    c.setFont("Helvetica-Bold", 14)
    c.drawRightString(x_text, y, "NANOSPECT AI")
    y -= 14
    c.drawRightString(x_text, y, "INSPECTIONS")

    c.setFont("Helvetica", 10)
    y -= 16
    c.drawRightString(x_text, y, "(870) 235-4000")
    y -= 12
    c.drawRightString(x_text, y, "Southern Arkansas University")
    y -= 12
    c.drawRightString(x_text, y, "Website Coming Soon")

    # Property photo + center text
    photo_w = 4.5 * inch
    photo_h = 2.5 * inch
    x = (width - photo_w) / 2
    y_img = height - 2.8 * inch - photo_h

    image = input("Enter Building Image (Path): ")
    c.drawImage(image, x, y_img, width=photo_w, height=photo_h, preserveAspectRatio=True, mask="auto")

    # Address & Building Type
    y_text = y_img - 24
    c.setFont("Helvetica", 14)
    buildingType = input("Enter The Type of Building: ")
    upperBuildingType = buildingType.upper()
    c.drawCentredString(width / 2, y_text, "NANOSPECT AI - " + upperBuildingType + " INSPECTION")

    c.setFont("Helvetica", 12)
    y_text -= 18
    c.drawCentredString(width / 2, y_text, "100 E University")
    y_text -= 14
    c.drawCentredString(width / 2, y_text, "Magnolia, AR 71753")

    c.setFont("Helvetica", 10)
    y_text -= 18
    building = input("Enter Building Name: ").strip() or ""
    c.drawCentredString(width / 2, y_text, building)
    y_text -= 12
    date = input("Enter The Date: ").strip() or ""
    c.drawCentredString(width / 2, y_text, date)

    # Divider line
    y_top = 3.5 * inch
    c.setLineWidth(0.5)
    c.line(margin, y_top, width - margin, y_top)

    # Left/middle/right member photos & info
    headshot = 1.5 * inch
    x_left = margin + 0.2 * inch
    y_image = y_top - 0.4 * inch - headshot

    # John photo 
    c.drawImage("John.jpg", x_left, y_image, width=headshot, height=headshot, mask="auto")


    cx_left = x_left + headshot / 2
    y_t = y_image - 14
    c.setFont("Helvetica", 9)
    c.drawCentredString(cx_left, y_t, "Group Member")
    y_t -= 12
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(cx_left, y_t, "John-Thomas Daigle")
    y_t -= 12
    c.setFont("Helvetica", 8)
    c.drawCentredString(cx_left, y_t, "JWDaigle5139@muleriders.saumag.edu")

    # Josh photo
    x_mid = (width - headshot) / 2
    c.drawImage("Josh.png", x_mid, y_image, width=headshot, height=headshot, mask="auto")

    cx_mid = x_mid + headshot / 2
    y_mid = y_image - 14
    c.setFont("Helvetica", 9)
    c.drawCentredString(cx_mid, y_mid, "Group Member")
    y_mid -= 12
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(cx_mid, y_mid, "Joshua Dumas")
    y_mid -= 12
    c.setFont("Helvetica", 8)
    c.drawCentredString(cx_mid, y_mid, "JPDumas4893@muleriders.saumag.edu")

    # Jackson Photo
    x_right = width - margin - 0.2 * inch - headshot
    c.drawImage("jackson_upright.jpg", x_right, y_image, width=headshot, height=headshot, mask="auto")

    cx_right = x_right + headshot / 2
    y_t = y_image - 14
    c.setFont("Helvetica", 9)
    c.drawCentredString(cx_right, y_t, "Group Member")
    y_t -= 12
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(cx_right, y_t, "Jackson Boyd")
    y_t -= 12
    c.setFont("Helvetica", 8)
    c.drawCentredString(cx_right, y_t, "JEBoyd4422@muleriders.saumag.edu")

    c.showPage()
    c.save()
    logger.info("Saved cover PDF to: %s", output_path)


#Image extraction and API Processing
def extract_image_caption_pairs(doc_path):
    """Extract images and guessed captions from DOCX (returns list of dicts)."""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Input DOCX not found: {doc_path}")

    os.makedirs(IMAGE_TEMP_DIR, exist_ok=True)
    document = Document(doc_path)

    pairs = []
    image_counter = 0
    paragraphs = document.paragraphs
    total_paragraphs = len(paragraphs)
    logger.info("Extracting images and captions from %s", doc_path)

    for i, paragraph in enumerate(paragraphs):
        runs_with_images = []
        for run in paragraph.runs:
            blips = run._element.xpath(".//a:blip")
            if blips:
                runs_with_images.append((run, blips))

        if not runs_with_images:
            continue

        # Guess caption, next non-empty paragraph
        caption_text = guess_caption(paragraphs, i, max_lookahead=3)
        # Extrack each image found
        for run, blips in runs_with_images:
            for blip in blips:
                r_id = blip.get(qn("r:embed"))
                if not r_id:
                    continue

                image_part = document.part.related_parts.get(r_id)
                if image_part is None:
                    continue
                # File type
                ext = os.path.splitext(image_part.partname)[1].lower()
                if not ext:
                    ext = ".jpeg"

                image_counter += 1
                image_name = f"image_{image_counter}{ext}"
                image_path = os.path.join(IMAGE_TEMP_DIR, image_name)
                with open(image_path, "wb") as f:
                    f.write(image_part.blob)

                logger.info("Extracted %s (paragraph %d/%d) with caption: %s",
                            image_name, i + 1, total_paragraphs, caption_text or "[no caption found]")

                pairs.append({
                    "image_path": image_path,
                    "image_name": image_name,
                    "context": caption_text,
                })

    logger.info("Total image–caption pairs extracted: %d", len(pairs))
    return pairs


def guess_caption(paragraphs, image_par_index, max_lookahead=3):
    """
    Attempts to guess a caption by looking at the paragraphs immediately
    following the image. Priority:
      1. Paragraphs using 'caption' style.
      2. First non-empty paragraph.
    """
    best_candidate = ""
    for j in range(image_par_index + 1, min(len(paragraphs), image_par_index + 1 + max_lookahead)):
        p = paragraphs[j]
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = p.style.name if p.style is not None else ""
        is_caption_style = "caption" in style_name.lower()
        if is_caption_style:
            return text
        if not best_candidate:
            best_candidate = text
    return best_candidate if best_candidate else ""

# AI Title Generation
def generate_image_title(image_path, context_text):
    img = Image.open(image_path)
    context_text = context_text or "(No inspection notes provided.)"

    prompt = f"""
You are analyzing an image from a property inspection report.

--- Inspector Notes ---
{context_text}
--- End Notes ---

Based on the image and the inspector's notes, generate a SHORT, DESCRIPTIVE TITLE (2-5 words) that identifies the main issue or defect shown.

Rules:
- Keep it SHORT (2-5 words maximum)
- Be SPECIFIC to what you see
- Use professional inspection terminology
- Do NOT use phrases like "Figure" or "Image of"
- Focus on the DEFECT or ISSUE, not just the location

Return ONLY the title, nothing else.
"""
    response = model.generate_content([prompt, img])
    title = (response.text or "").strip()
    title = title.strip('"').strip("'").strip()
    return title if title else "Inspection Finding"

#Inspection description paragraph generator
def generate_inspection_description(image_path, context_text):
    img = Image.open(image_path)
    context_text = context_text or "(No inspection notes provided.)"

    prompt = f"""
You are generating text for a professional property inspection report.
You are given:
1. An image of the issue.
2. Inspector notes taken between the images in the document.

Use BOTH to produce a clear, professional inspection-quality description.

--- Inspector Notes ---
{context_text}
--- End Notes ---

Write a detailed paragraph that includes:
- What the image shows (condition, material, identifiable defects)
- How the notes relate to the defect
- Likely cause (if observable from image and text)
- Potential implications (water damage, structural issues, cosmetic issue, etc.)
- Recommended corrective action (repair, replacement, caulking, sealing, monitoring, etc.)

Rules:
- Do NOT mention "caption."
- Do NOT mention that this text came from notes.
- Do NOT speculate beyond what the image + text reasonably support.
- Write in a professional, objective tone suitable for an insurance-reviewable inspection report.
"""
    response = model.generate_content([prompt, img])
    return (response.text or "")

#Create the DOCX containing image, title, and generated description(Work in progress)
def build_output_docx(pairs, output_path):
    doc = Document()
    doc.add_heading("Inspection Images and Descriptions", level=1)
    for idx, item in enumerate(pairs, start=1):
        image_path = item["image_path"]
        description = item.get("description", "")
        title = item.get("title", "Inspection Finding")

        doc.add_heading(f"Figure {idx}: {title}", level=2)

        try:
            doc.add_picture(image_path, width=Inches(4))
        except Exception as e:
            logger.warning("Could not insert image %s: %s", image_path, e)
            doc.add_paragraph(f"[Could not display image: {e}]")

        if description:
            p_desc = doc.add_paragraph()
            run_label = p_desc.add_run("Inspection Description: ")
            run_label.bold = True
            p_desc.add_run(description)
        else:
            doc.add_paragraph("Inspection Description: [not generated]")

        doc.add_paragraph("")  # spacing

    doc.save(output_path)
    logger.info("Saved output DOCX to: %s", output_path)


def convert_docx_to_pdf(input_path, output_path):
    # Try docx2pdf if installed
    if DOCX2PDF_AVAILABLE:
        try:
            logger.info("Attempting conversion using docx2pdf (requires Word on Windows)...")
            docx2pdf_convert(input_path, output_path)
            logger.info("Saved PDF to: %s", output_path)
            return True
        except Exception as e:
            logger.warning("docx2pdf conversion failed: %s", e)

    # Fallback to LibreOffice 
    try:
        logger.info("Attempting conversion using LibreOffice (soffice)...")
        # Build command. Use --headless to avoid GUI. Output to a temporary dir then move file.
        out_dir = Path(output_path).resolve().parent
        cmd = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(Path(input_path).resolve()),
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        # LibreOffice names the output same as input with .pdf extension in out_dir
        generated_pdf = out_dir / (Path(input_path).stem + ".pdf")
        if generated_pdf.exists():
            generated_pdf.rename(output_path)
            logger.info("Saved PDF to: %s", output_path)
            return True
        else:
            logger.warning("LibreOffice conversion did not produce expected file: %s", generated_pdf)
    except Exception as e:
        logger.warning("LibreOffice conversion failed or not available: %s", e)

    logger.error("DOCX to PDF conversion failed. Please install Word (with docx2pdf) or LibreOffice (soffice).")
    return False


def merge_pdfs(pdf_paths, output_path):
    """Merge PDFs in order using pypdf."""
    writer = PdfWriter()
    for p in pdf_paths:
        if not os.path.exists(p):
            logger.error("Cannot merge; PDF not found: %s", p)
            raise FileNotFoundError(p)
        reader = PdfReader(p)
        for page in reader.pages:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)
    logger.info("Merged PDFs into: %s", output_path)

def main():
    input_docx = sys.argv[1] if len(sys.argv) > 1 else INPUT_DOCX
    if not os.path.exists(input_docx):
        logger.error("Input DOCX not found: %s", input_docx)
        return

    #Build DOCX page cover(Work In Progress)
    cover_doc = Document()

    # Ask the user for inputs
    print("=== Cover Page Inputs ===")
    logo_path = input("Enter path to SAU logo image (saulogo.jpg recommended): ").strip() or "saulogo.jpg"
    image_path = input("Enter Image Name (path) for property photo: ").strip()
    buildingType = input("Enter The Type of Building: ").strip().upper()
    buildingName = input("Enter Building Name: ").strip()
    date = input("Enter The Date: ").strip()

    # Build the DOCX cover
    try:
        cover_doc.add_picture(logo_path, width=Inches(1.3))
    except:
        pass

    cover_doc.add_heading("NANOSPECT AI INSPECTIONS", level=1)
    cover_doc.add_paragraph("(870) 235-4000\nSouthern Arkansas University\nWebsite Coming Soon")

    if image_path:
        try:
            cover_doc.add_picture(image_path, width=Inches(4.5))
        except:
            pass

    cover_doc.add_heading(f"NANOSPECT AI - {buildingType} INSPECTION", level=2)
    cover_doc.add_paragraph("100 E University")
    cover_doc.add_paragraph("Magnolia, AR 71753")
    cover_doc.add_paragraph(buildingName)
    cover_doc.add_paragraph(date)

    cover_doc.add_page_break()

    # Extract images
    pairs = extract_image_caption_pairs(input_docx)

    # Generate titles + descriptions
    for item in pairs:
        try:
            item["title"] = generate_image_title(item["image_path"], item["context"])
        except:
            item["title"] = "Inspection Finding"

        try:
            item["description"] = generate_inspection_description(item["image_path"], item["context"])
        except:
            item["description"] = ""

    #Append inspection results to cover_doc

    cover_doc.add_heading("Inspection Images and Descriptions", level=1)

    for idx, item in enumerate(pairs, start=1):
        image_path = item["image_path"]
        title = item["title"]
        description = item["description"]

        h = cover_doc.add_paragraph(style="Heading 2")
        h.add_run(f"Figure {idx}: {title}")

        try:
            cover_doc.add_picture(image_path, width=Inches(4))
        except:
            cover_doc.add_paragraph("[Could not display image]")

        p = cover_doc.add_paragraph()
        b = p.add_run("Inspection Description: ")
        b.bold = True
        p.add_run(description)

        cover_doc.add_paragraph("")

    #SAVE THE FINAL DOCX
    FINAL_DOCX = "final_report.docx"
    cover_doc.save(FINAL_DOCX)
    print(f"\nAll done — final DOCX created: {FINAL_DOCX}")

#Main
def main():
    input_docx = sys.argv[1] if len(sys.argv) > 1 else INPUT_DOCX
    if not os.path.exists(input_docx):
        logger.error("Input DOCX not found: %s", input_docx)
        return

    #Create interactive cover page for PDF
    create_cover(COVER_PDF)

    #Extract images + captions
    pairs = extract_image_caption_pairs(input_docx)
    if not pairs:
        logger.warning("No images found in the document. Building empty inspection DOCX.")
        # Still build an empty docx
        build_output_docx([], OUTPUT_DOCX)
        # Convert empty docx to PDF
        converted = convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)
        if not converted:
            logger.error("Could not produce inspection PDF. Aborting merge.")
            return
        # Merge cover + inspection 
        merge_pdfs([COVER_PDF, OUTPUT_PDF], FINAL_PDF)
        logger.info("Final PDF saved to: %s", FINAL_PDF)
        return

    #Generate titles for each image
    logger.info("Generating titles for %d images...", len(pairs))
    for item in pairs:
        try:
            item["title"] = generate_image_title(item["image_path"], item["context"])
            logger.info("Title: %s -> %s", item["image_name"], item["title"])
        except Exception as e:
            logger.error("Error generating title for %s: %s", item["image_name"], e)
            item["title"] = "Inspection Finding"

    #Generate descriptions
    logger.info("Generating descriptions for %d images...", len(pairs))
    for item in pairs:
        try:
            item["description"] = generate_inspection_description(item["image_path"], item["context"])
            logger.info("Description generated for %s", item["image_name"])
        except Exception as e:
            logger.error("Error generating description for %s: %s", item["image_name"], e)
            item["description"] = ""

    #Build output DOCX with results
    build_output_docx(pairs, OUTPUT_DOCX)

    #Convert DOCX to PDF
    converted = convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)
    if not converted:
        logger.error("Failed to convert inspection DOCX to PDF. Aborting merge.")
        return

    # 7) Merge cover + inspection into final PDF
    merge_pdfs([COVER_PDF, OUTPUT_PDF], FINAL_PDF)
    logger.info("Final merged PDF created: %s", FINAL_PDF)
    print(f"\nAll done — final PDF: {FINAL_PDF}")


if __name__ == "__main__":
    main()