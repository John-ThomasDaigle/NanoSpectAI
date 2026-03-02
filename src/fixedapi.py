import logging
import os
import subprocess
from pathlib import Path

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from PIL import Image

try:
    from pypdf import PdfReader, PdfWriter
except Exception:
    from PyPDF2 import PdfReader, PdfWriter

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

import google.generativeai as genai


# ---------------------------
# CONFIG
# ---------------------------

MODEL_NAME = "gemini-2.5-pro"
IMAGE_TEMP_DIR = "extracted_images"
OUTPUT_DOCX = "inspection_output.docx"
OUTPUT_PDF = "inspection_output.pdf"
COVER_PDF = "cover.pdf"
FINAL_PDF = "final_report.pdf"

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise RuntimeError("GEMINI_API_KEY not set.")
genai.configure(api_key=api_key)
model = genai.GenerativeModel(MODEL_NAME)


# ---------------------------
# COVER GENERATION (UI READY)
# ---------------------------

def create_cover(output_path, building_type, building_name, inspection_date, building_image_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, height - 1.5 * inch,
                        f"NANOSPECT AI - {building_type.upper()} INSPECTION")

    c.setFont("Helvetica", 12)
    c.drawCentredString(width / 2, height - 1.9 * inch, building_name)
    c.drawCentredString(width / 2, height - 2.2 * inch, inspection_date)

    if building_image_path and os.path.exists(building_image_path):
        c.drawImage(building_image_path,
                    1.5 * inch,
                    height - 5 * inch,
                    width=4 * inch,
                    preserveAspectRatio=True,
                    mask="auto")

    c.showPage()
    c.save()


# ---------------------------
# IMAGE EXTRACTION
# ---------------------------

def extract_image_caption_pairs(doc_path):
    if not os.path.exists(doc_path):
        raise FileNotFoundError(doc_path)

    os.makedirs(IMAGE_TEMP_DIR, exist_ok=True)
    document = Document(doc_path)

    pairs = []
    image_counter = 0

    for i, paragraph in enumerate(document.paragraphs):
        for run in paragraph.runs:
            blips = run._element.xpath(".//a:blip")
            if not blips:
                continue

            for blip in blips:
                r_id = blip.get(qn("r:embed"))
                image_part = document.part.related_parts.get(r_id)
                if not image_part:
                    continue

                ext = os.path.splitext(image_part.partname)[1] or ".jpeg"
                image_counter += 1
                image_path = os.path.join(IMAGE_TEMP_DIR, f"image_{image_counter}{ext}")

                with open(image_path, "wb") as f:
                    f.write(image_part.blob)

                pairs.append({
                    "image_path": image_path,
                    "context": ""
                })

    return pairs


# ---------------------------
# AI GENERATION
# ---------------------------

def generate_image_title(image_path):
    img = Image.open(image_path)
    prompt = "Generate a 2-5 word professional inspection title for this image."
    response = model.generate_content([prompt, img])
    return (response.text or "").strip() or "Inspection Finding"


def generate_inspection_description(image_path):
    img = Image.open(image_path)
    prompt = "Generate a professional inspection description paragraph for this image."
    response = model.generate_content([prompt, img])
    return (response.text or "").strip()


# ---------------------------
# DOCX BUILD
# ---------------------------

def build_output_docx(pairs, output_path):
    doc = Document()
    doc.add_heading("Inspection Images and Descriptions", level=1)

    for idx, item in enumerate(pairs, start=1):
        doc.add_heading(f"Figure {idx}: {item['title']}", level=2)
        doc.add_picture(item["image_path"], width=Inches(4))

        p = doc.add_paragraph()
        b = p.add_run("Inspection Description: ")
        b.bold = True
        p.add_run(item["description"])

    doc.save(output_path)


# ---------------------------
# PDF CONVERSION
# ---------------------------

def convert_docx_to_pdf(input_path, output_path):
    if DOCX2PDF_AVAILABLE:
        docx2pdf_convert(input_path, output_path)
        return True

    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            input_path
        ], check=True)
        return True
    except Exception:
        return False


def merge_pdfs(paths, output_path):
    writer = PdfWriter()
    for p in paths:
        reader = PdfReader(p)
        for page in reader.pages:
            writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)


# ---------------------------
# MAIN FUNCTION FOR STREAMLIT
# ---------------------------

def generate_inspection_report(
    input_docx_path,
    building_type,
    building_name=None,
    inspection_date=None,
    building_image_path=None
):
    create_cover(COVER_PDF,
                 building_type,
                 building_name,
                 inspection_date,
                 building_image_path)

    pairs = extract_image_caption_pairs(input_docx_path)

    for item in pairs:
        item["title"] = generate_image_title(item["image_path"])
        item["description"] = generate_inspection_description(item["image_path"])

    build_output_docx(pairs, OUTPUT_DOCX)

    if not convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF):
        raise RuntimeError("PDF conversion failed.")

    merge_pdfs([COVER_PDF, OUTPUT_PDF], FINAL_PDF)

    return FINAL_PDF