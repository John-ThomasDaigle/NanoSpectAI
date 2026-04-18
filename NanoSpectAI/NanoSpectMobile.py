import torch
import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wavfile
import whisper
import time
import threading
import os
import re
from PIL import Image
from docx import Document
from docx.shared import Inches
from datetime import timedelta, datetime
import warnings
warnings.filterwarnings("ignore", message="FP16 is not supported on CPU; using FP32 instead")

# Root directory for saving audio and output files

BASE_DIR    = '/home/student/audio-transcription-using-openai-whisper'

# Settings
SAMPLE_RATE = 16000
FILENAME    = os.path.join(BASE_DIR, "recording.wav")
NAME        = "Jetson Original Report Mobile"

# Folder where Bluetooth file transfers are received
FOLDERNAME  = os.path.expanduser("~/Downloads")

# Stores images
captured_images = []

# Timestamps set when recording begins
record_start          = None
record_start_datetime = None

# Event flag used to signal all threads to stop recording
stop_recording = threading.Event()

# Matches common image file extensions
IMAGE_PATTERN = re.compile(r'\.(jpe?g|png)$', re.IGNORECASE)

# Matches any file named STOP with any or no extension
STOP_PATTERN  = re.compile(r'^STOP(\..+)?$', re.IGNORECASE) 

def countdown(seconds):
    # Prints a countdown to let the user prepare before recording starts
    print("Get ready to record")
    for i in range(seconds, 0, -1):
        print(i)
        time.sleep(1)
    print("Recording started!")


def get_image_datetime(filepath):
    """
    Read capture time from EXIF metadata.
    Falls back to file modification time if EXIF is unavailable.
    """
    try:
        img      = Image.open(filepath)
        exif_raw = img._getexif()
        if exif_raw:
            for tag_id in (36867, 36868, 306):   # DateTimeOriginal, DateTimeDigitized, DateTime
                value = exif_raw.get(tag_id)
                if value:
                    return datetime.strptime(value, "%Y:%m:%d %H:%M:%S")
    except Exception:
        pass
    return datetime.fromtimestamp(os.path.getmtime(filepath))

def record_audio():
    # Continuously records audio from the microphone until stop_recording is set
    audio_buffer = []

    def callback(indata, frames, time_info, status):
        # Appends each incoming audio chunk to the buffer, stops when flagged
        if stop_recording.is_set():
            raise sd.CallbackStop()
        audio_buffer.append(indata.copy())

    with sd.InputStream(samplerate=SAMPLE_RATE, channels=1,
                        dtype='int16', callback=callback):
        while not stop_recording.is_set():
            time.sleep(0.1)

    # Combine all chunks and save to a WAV file
    audio_data = np.concatenate(audio_buffer, axis=0)
    wavfile.write(FILENAME, SAMPLE_RATE, audio_data)
    print(f"Recording saved as {FILENAME}")

def watch_for_bluetooth_images():
    # Watches the Bluetooth download folder for new files
    # Logs received images and stops recording when a STOP file is detected
    print("\nWaiting for images via Bluetooth...")
    print(f'Send images from your phone. Send a file named "STOP" to finish.\n')

    seen_files = set(os.listdir(FOLDERNAME))

    while not stop_recording.is_set():
        current_files = set(os.listdir(FOLDERNAME))
        new_files     = current_files - seen_files

        for filename in sorted(new_files):
            full_path = os.path.join(FOLDERNAME, filename)

            # A file named STOP signals the end of the recording session
            if STOP_PATTERN.match(filename):
                print(f'\n  Received stop signal ("{filename}"). Ending recording...')
                stop_recording.set()
                try:
                    os.remove(full_path)
                except OSError:
                    pass
                break

            # Log the image with its capture time if it's a photo
            if IMAGE_PATTERN.search(filename):
                time.sleep(0.3)   # allow file to finish writing
                img_time = get_image_datetime(full_path)
                captured_images.append((full_path, img_time))
                print(f"  Received: {filename}  "
                      f"(taken {img_time:%a %d %b %Y %I:%M:%S %p})")

        seen_files = current_files
        time.sleep(0.5)

    print(f"\nRecording stopped. {len(captured_images)} image(s) logged.")

def save_transcription_with_timestamps(transcription, captured_images, output_file=None):
    # Builds a Word document that pairs each received image with the transcript
    # segments spoken while that image was the most recently captured one
    if output_file is None:
        output_file = os.path.join(BASE_DIR, f"{NAME}.docx")
    
    doc = Document()

    # Sort images by actual capture time (EXIF / mtime)
    captured_images.sort(key=lambda x: x[1])

    # Create a bucket for each image to hold its associated transcript segments
    groups = {img: [] for img, _ in captured_images}

    for seg in transcription.get("segments", []):
        # Find the midpoint of the segment and convert it to a real-world timestamp
        seg_mid    = (seg["start"] + seg["end"]) / 2.0
        seg_mid_dt = record_start_datetime + timedelta(seconds=seg_mid)

        # Assign the segment to the last image captured before the segment's midpoint
        target_image = None
        for img_path, img_dt in captured_images:
            if seg_mid_dt >= img_dt:
                target_image = img_path
            else:
                break

        if target_image:
            groups[target_image].append(seg)

    # Write each image section and its transcript
    for img_path, img_dt in captured_images:
        img_filename = os.path.basename(img_path)

        doc.add_heading(
            f"{img_filename} captured at {img_dt.strftime('%a %d %b %Y %I:%M:%S %p')}",
            level=2
        )

        try:
            doc.add_picture(img_path, width=Inches(3))
        except Exception as e:
            doc.add_paragraph(f"(Error loading image: {e})")

        # Add each segment as a timestamped line of text
        for seg in groups[img_path]:
            start_dt = record_start_datetime + timedelta(seconds=seg["start"])
            end_dt   = record_start_datetime + timedelta(seconds=seg["end"])
            line = (f"[{start_dt.strftime('%a %d %b %Y %I:%M:%S %p')} → "
                    f"{end_dt.strftime('%I:%M:%S %p')}]  {seg['text']}")
            doc.add_paragraph(line)

    doc.save(output_file)
    print(f"Saved: {output_file}")

# Load Whisper model 
model = whisper.load_model("base").to(
    "cuda" if torch.cuda.is_available() else "cpu"
)

def main():
    countdown(3)

    global record_start, record_start_datetime
    record_start          = time.time()
    record_start_datetime = datetime.now()

    # Start audio recording in a background thread
    audio_thread = threading.Thread(target=record_audio)
    audio_thread.start()

    # Watch for Bluetooth images (blocks until STOP file received)
    watch_for_bluetooth_images()

    # Wait for audio to finish saving
    audio_thread.join()

    # Transcribe
    print("Transcribing...")
    result       = model.transcribe(FILENAME)
    english_text = result["text"]
    print(f"Full Transcription: {english_text}")

    # Save Word document
    print("Saving transcription with timestamps...")
    save_transcription_with_timestamps(result, captured_images)
    print(f"Transcription saved as '{NAME}.docx'.")


if __name__ == "__main__":
    main()