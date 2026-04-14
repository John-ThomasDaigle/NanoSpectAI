import torch
import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wavfile
import whisper
import time
import threading
import cv2
import os
import re
from docx import Document
from docx.shared import Inches
from datetime import timedelta, datetime
import warnings
warnings.filterwarnings("ignore", message="FP16 is not supported on CPU; using FP32 instead")

BASE_DIR   = '/home/student/audio-transcription-using-openai-whisper'

# Settings
SAMPLE_RATE = 16000  # 16 khz
FILENAME = os.path.join(BASE_DIR, "recording.wav")
NAME = "Jetson Original Report Camera"
FOLDERNAME = os.path.join(BASE_DIR,"captured_images")
captured_images = []
os.makedirs(FOLDERNAME, exist_ok=True)

# These will be set when recording begins
record_start = None
record_start_datetime = None

# Function to display countdown before recording starts
def countdown(seconds):
    print("Get ready to record")
    for i in range(seconds, 0, -1):
        print(i, "")
        time.sleep(1)
    print("Recording started!")

# Global stop flag
stop_recording = threading.Event()

# Function to record audio until stop_recoding is set
def record_audio():
    audio_buffer = []
    
    def callback(indata, frames, time_info, status):
        if stop_recording.is_set():
            raise sd.CallbackStop()
        audio_buffer.append(indata.copy())
    
    with sd.InputStream(samplerate=SAMPLE_RATE, channels=1, dtype='int16', callback=callback):
        while not stop_recording.is_set():
            time.sleep(0.1)

    audio_data = np.concatenate(audio_buffer, axis=0)
    wavfile.write(FILENAME, SAMPLE_RATE, audio_data)
    print(f"Recording saved as {FILENAME}")

# Initialize Whisper model
model = whisper.load_model("base").to("cuda" if torch.cuda.is_available() else "cpu")

# Camera Capture loop that saves images and sets stop_recording on ESC
def capture_camera():
    # Find the next image number automatically
    existing_files = [f for f in os.listdir(FOLDERNAME) if f.startswith("image_") and f.endswith(".jpeg")]
    # Extract numbers from filenames like image_1.jpeg
    numbers = []
    for f in existing_files:
        match = re.search(r'image_(\d+)\.jpeg', f)
        if match:
            numbers.append(int(match.group(1)))

    # Start counting from the next number, or 1 if no files exist
    img_count = max(numbers) if numbers else 0

    # Open camera
    cap = cv2.VideoCapture(0)
    if not cap.isOpened():
        print("Error: Could not open camera.")
        exit()

    print("Press SPACE to capture an image.")
    print("Press ESC to exit.")

    while True:
        ret, frame = cap.read()
        if not ret:
            print("Failed to grab frame.")
            break

        cv2.imshow("USB Camera", frame)
        key = cv2.waitKey(1)

        # Closes program and allows the transcription to start
        if key == 27:  # ESC key
            print("Exiting...")
            stop_recording.set() # Stops Whisper recording
            break
        # Saves images 
        elif key == 32:  # SPACE key
            img_count += 1
            filename = os.path.join(FOLDERNAME, f"image_{img_count}.jpeg")
            cv2.imwrite(filename, frame)
            image_time = datetime.now()
            captured_images.append((filename, image_time))
            print(f"Saved: {filename} at {image_time}")

    cap.release()
    cv2.destroyAllWindows()

# Save transcripion and images into a Word file with timestamps
def save_transcription_with_timestamps(transcription, captured_images, output_file=None):
    if output_file is None:
        output_file = os.path.join(BASE_DIR, f"{NAME}.docx")
    # Create a Word document that shows each image with its timestamp and speech taht happened after it
    doc = Document()

    # Sort captured images in time order
    captured_images.sort(key=lambda x:int(re.search(r'image_(\d+)', x[0]).group(1))) 

    # Prepare segment grouping
    groups = {img: [] for img, _ in captured_images}
    before_first_image = []

    # Match each audio segment to the appropriate image
    for seg in transcription.get("segments", []):
        seg_start = seg["start"]
        seg_end = seg["end"]
        seg_mid = (seg_start + seg_end) / 2.0  # midpoint of speech

        seg_mid_dt = record_start_datetime + timedelta(seconds=seg_mid)

        target_image = None

        # Find last image that occurred before the speech segment
        for img_path, img_dt in captured_images:
            if seg_mid_dt >= img_dt:
                target_image = img_path
            else:
                break

        # If we found an image, add segment to that image's group, else store as before_first_image
        if target_image:
            groups[target_image].append(seg)
        else:
            before_first_image.append(seg)

    # Write each image and its segments
    for img_path, img_dt in captured_images:
        #Get the exact saved filename, like "image_7.jpeg"
        img_filename = os.path.basename(img_path)

        # Image header line with name and capture time
        doc.add_heading(
            f"{img_filename} captured at {img_dt.strftime('%a %d %b %Y %I:%M:%S %p')}",
            level=2
        )

        # Insert image into docx
        try:
            doc.add_picture(img_path, width=Inches(3))
        except Exception as e:
            doc.add_paragraph(f"(Error loading image: {e})")

        # Add all speech assigned to this image with timestamps
        for seg in groups[img_path]:
            start_dt = record_start_datetime + timedelta(seconds=seg["start"])
            end_dt = record_start_datetime + timedelta(seconds=seg["end"])
            line = f"[{start_dt.strftime('%a %d %b %Y %I:%M:%S %p')} → {end_dt.strftime('%I:%M:%S %p')}]  {seg['text']}"
            doc.add_paragraph(line)

    # Save output .docx
    doc.save(output_file)
    print(f"Saved: {output_file}")

# Main function to run countdown, audio, camera, and transcription
def main():
    #Set coundown timer before recording
    countdown(3)

    #Record both numeric and datetime start times
    global record_start, record_start_datetime
    record_start = time.time()
    record_start_datetime = datetime.now()

    # Start audio recording in a separate thread
    audio_thread = threading.Thread(target=record_audio)
    audio_thread.start()

    # Run camera capture in the main thread
    capture_camera()

    # Wait for audio recording to finish
    audio_thread.join()

    # Transcribe audio file with Whisper
    print("Transcribing...")
    result = model.transcribe(FILENAME)

    # Display transcription in console
    english_text = result["text"]
    print(f"Full Transcription: {english_text}")

    # Save transcription with timestamps to a Word file
    print("Saving transcription with timestamps...")
    save_transcription_with_timestamps(result, captured_images)
    print(f"Transcription saved as '{NAME}.docx'.")

if __name__ == "__main__":
    main()
