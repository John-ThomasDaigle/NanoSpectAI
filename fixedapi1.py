import logging
import os
import subprocess
from pathlib import Path

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from PIL import Image

try:
    from pypdf import PdfReader, PdfWriter
except Exception:
    from PyPDF2 import PdfReader, PdfWriter

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

import google.generativeai as genai


# ---------------------------
# CONFIG
# ---------------------------

MODEL_NAME = "gemini-2.5-pro"
IMAGE_TEMP_DIR = "extracted_images"
OUTPUT_DOCX = "inspection_output.docx"
OUTPUT_PDF = "inspection_output.pdf"
COVER_PDF = "cover.pdf"
FINAL_PDF = "final_report.pdf"

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise RuntimeError("GEMINI_API_KEY not set.")
genai.configure(api_key=api_key)
model = genai.GenerativeModel(MODEL_NAME)


# ---------------------------
# COVER GENERATION (UI READY)
# ---------------------------

def create_cover(output_path, building_type, building_name, inspection_date, building_image_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, height - 1.5 * inch,
                        f"NANOSPECT AI - {building_type.upper()} INSPECTION")

    c.setFont("Helvetica", 12)
    c.drawCentredString(width / 2, height - 1.9 * inch, building_name)
    c.drawCentredString(width / 2, height - 2.2 * inch, inspection_date)

    if building_image_path and os.path.exists(building_image_path):
        c.drawImage(building_image_path,
                    1.5 * inch,
                    height - 5 * inch,
                    width=4 * inch,
                    preserveAspectRatio=True,
                    mask="auto")

    c.showPage()
    c.save()


#Image extraction and API Processing
def extract_image_caption_pairs(doc_path):
    """Extract images and guessed captions from DOCX (returns list of dicts)."""
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Input DOCX not found: {doc_path}")

    os.makedirs(IMAGE_TEMP_DIR, exist_ok=True)
    document = Document(doc_path)

    pairs = []
    image_counter = 0
    paragraphs = document.paragraphs
    total_paragraphs = len(paragraphs)
    logger.info("Extracting images and captions from %s", doc_path)

    for i, paragraph in enumerate(paragraphs):
        runs_with_images = []
        for run in paragraph.runs:
            blips = run._element.xpath(".//a:blip")
            if blips:
                runs_with_images.append((run, blips))

        if not runs_with_images:
            continue

        # Guess caption, next non-empty paragraph
        caption_text = guess_caption(paragraphs, i, max_lookahead=3)
        # Extrack each image found
        for run, blips in runs_with_images:
            for blip in blips:
                r_id = blip.get(qn("r:embed"))
                if not r_id:
                    continue

                image_part = document.part.related_parts.get(r_id)
                if image_part is None:
                    continue
                # File type
                ext = os.path.splitext(image_part.partname)[1].lower()
                if not ext:
                    ext = ".jpeg"

                image_counter += 1
                image_name = f"image_{image_counter}{ext}"
                image_path = os.path.join(IMAGE_TEMP_DIR, image_name)
                with open(image_path, "wb") as f:
                    f.write(image_part.blob)

                logger.info("Extracted %s (paragraph %d/%d) with caption: %s",
                            image_name, i + 1, total_paragraphs, caption_text or "[no caption found]")

                pairs.append({
                    "image_path": image_path,
                    "image_name": image_name,
                    "context": caption_text,
                })

    logger.info("Total image–caption pairs extracted: %d", len(pairs))
    return pairs


def guess_caption(paragraphs, image_par_index, max_lookahead=3):
    """
    Attempts to guess a caption by looking at the paragraphs immediately
    following the image. Priority:
      1. Paragraphs using 'caption' style.
      2. First non-empty paragraph.
    """
    best_candidate = ""
    for j in range(image_par_index + 1, min(len(paragraphs), image_par_index + 1 + max_lookahead)):
        p = paragraphs[j]
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = p.style.name if p.style is not None else ""
        is_caption_style = "caption" in style_name.lower()
        if is_caption_style:
            return text
        if not best_candidate:
            best_candidate = text
    return best_candidate if best_candidate else ""

# AI Title Generation
def generate_image_title(image_path, context_text):
    img = Image.open(image_path)
    context_text = context_text or "(No inspection notes provided.)"

    prompt = f"""
You are analyzing an image from a property inspection report.

--- Inspector Notes ---
{context_text}
--- End Notes ---

Based on the image and the inspector's notes, generate a SHORT, DESCRIPTIVE TITLE (2-5 words) that identifies the main issue or defect shown.

Rules:
- Keep it SHORT (2-5 words maximum)
- Be SPECIFIC to what you see
- Use professional inspection terminology
- Do NOT use phrases like "Figure" or "Image of"
- Focus on the DEFECT or ISSUE, not just the location

Return ONLY the title, nothing else.
"""
    response = model.generate_content([prompt, img])
    title = (response.text or "").strip()
    title = title.strip('"').strip("'").strip()
    return title if title else "Inspection Finding"

#Inspection description paragraph generator
def generate_inspection_description(image_path, context_text):
    img = Image.open(image_path)
    context_text = context_text or "(No inspection notes provided.)"

    prompt = f"""
You are generating text for a professional property inspection report.
You are given:
1. An image of the issue.
2. Inspector notes taken between the images in the document.

Use BOTH to produce a clear, professional inspection-quality description.

--- Inspector Notes ---
{context_text}
--- End Notes ---

Write a detailed paragraph that includes:
- What the image shows (condition, material, identifiable defects)
- How the notes relate to the defect
- Likely cause (if observable from image and text)
- Potential implications (water damage, structural issues, cosmetic issue, etc.)
- Recommended corrective action (repair, replacement, caulking, sealing, monitoring, etc.)

Rules:
- Do NOT mention "caption."
- Do NOT mention that this text came from notes.
- Do NOT speculate beyond what the image + text reasonably support.
- Write in a professional, objective tone suitable for an insurance-reviewable inspection report.
"""
    response = model.generate_content([prompt, img])
    return (response.text or "")


# ---------------------------
# DOCX BUILD
# ---------------------------

def build_output_docx(pairs, output_path):
    doc = Document()
    doc.add_heading("Inspection Images and Descriptions", level=1)

    for idx, item in enumerate(pairs, start=1):
        doc.add_heading(f"Figure {idx}: {item['title']}", level=2)
        doc.add_picture(item["image_path"], width=Inches(4))

        p = doc.add_paragraph()
        b = p.add_run("Inspection Description: ")
        b.bold = True
        p.add_run(item["description"])

    doc.save(output_path)


# ---------------------------
# PDF CONVERSION
# ---------------------------

def convert_docx_to_pdf(input_path, output_path):
    if DOCX2PDF_AVAILABLE:
        docx2pdf_convert(input_path, output_path)
        return True

    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            input_path
        ], check=True)
        return True
    except Exception:
        return False


def merge_pdfs(paths, output_path):
    writer = PdfWriter()
    for p in paths:
        reader = PdfReader(p)
        for page in reader.pages:
            writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)


# ---------------------------
# MAIN FUNCTION FOR STREAMLIT
# ---------------------------

def generate_inspection_report(
    input_docx_path,
    building_type,
    building_name=None,
    inspection_date=None,
    building_image_path=None
):
    create_cover(COVER_PDF,
                 building_type,
                 building_name,
                 inspection_date,
                 building_image_path)

    pairs = extract_image_caption_pairs(input_docx_path)

    for item in pairs:
        item["title"] = generate_image_title(item["image_path"])
        item["description"] = generate_inspection_description(item["image_path"])

    build_output_docx(pairs, OUTPUT_DOCX)

    if not convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF):
        raise RuntimeError("PDF conversion failed.")

    merge_pdfs([COVER_PDF, OUTPUT_PDF], FINAL_PDF)

    return FINAL_PDF